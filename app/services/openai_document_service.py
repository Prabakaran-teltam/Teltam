import os
import re
import time
import base64
import hashlib
import logging
from django.conf import settings
from django.core.cache import cache
from openai import OpenAI, RateLimitError, APIConnectionError, APIStatusError

logger = logging.getLogger(__name__)

# Initialize the OpenAI Client
client = None

def get_openai_client():
    global client
    if not client:
        api_key = getattr(settings, 'OPENAI_API_KEY', None)
        if api_key:
            client = OpenAI(api_key=api_key)
    return client

def extract_text_from_image_with_openai(image_path):
    """
    Sends an image file to OpenAI's gpt-4o-mini Vision model to extract structured text content.
    """
    openai_client = get_openai_client()
    if not openai_client:
        raise ValueError("OpenAI API client is not configured. Please set OPENAI_API_KEY.")

    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Image not found at: {image_path}")

    # Read and encode image in base64
    with open(image_path, "rb") as image_file:
        base64_image = base64.b64encode(image_file.read()).decode("utf-8")

    ext = os.path.splitext(image_path)[1].lower().replace(".", "")
    mime_type = f"image/{ext}"
    if ext == "jpg" or ext == "jpeg":
        mime_type = "image/jpeg"
    elif ext == "png":
        mime_type = "image/png"

    response = openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an ultra-high precision Optical Character Recognition (OCR) model.\n"
                    "Your task is to extract EVERY SINGLE piece of text visible in the document image with 100% completeness and verbatim accuracy.\n\n"
                    "STRICT OCR RULES:\n"
                    "1. EXTRACT ALL VISIBLE TEXT: Include main body text, headers, subheadings, footers, page numbers, captions, table contents, cell text, bullet points, numbered lists, sidebars, logos/stamps with text, badge labels, and fine print.\n"
                    "2. DO NOT SUMMARIZE, SKIP, OR OMIT: Transcribe every single word, letter, number, punctuation mark, and symbol. Never skip text just because it is small, rotated, repeated, or formatted in columns/tables.\n"
                    "3. DO NOT TRANSLATE: Output the extracted text in its exact original language and script (e.g. Tamil, Hindi, English, Spanish, Chinese, Arabic, etc.).\n"
                    "4. PRESERVE LAYOUT: Preserve paragraphs, table grid structures (using pipe | separators), line breaks, and indentation as much as possible.\n"
                    "5. OUTPUT ONLY EXTRACTED TEXT: Do not add intro greetings, notes, explanations, markdown quotes, or metadata tags. Return strictly the transcribed text."
                )
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "text", 
                        "text": "Perform exhaustive 100% full-text OCR on this image. Extract every single word, header, footer, table entry, number, and note with absolute precision."
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{base64_image}",
                            "detail": "high"
                        }
                    }
                ]
            }
        ],
        max_tokens=4096,
        temperature=0.0
    )
    return response.choices[0].message.content.strip()

def extract_text_from_multiple_images(image_paths):
    """
    Extracts structured text from a batch/list of image files using OpenAI's gpt-4o-mini Vision model.
    Labels each image's text cleanly with header banners.
    """
    if not image_paths:
        return ""
        
    all_extracted = []
    for idx, img_path in enumerate(image_paths, 1):
        filename = os.path.basename(img_path)
        try:
            txt = extract_text_from_image_with_openai(img_path)
            if not txt.strip():
                txt = f"[Image {idx}: No readable text found]"
            all_extracted.append(f"--- Image {idx} ({filename}) ---\n{txt}")
        except Exception as err:
            logger.warning(f"Failed Vision OCR for image {filename}: {str(err)}")
            all_extracted.append(f"--- Image {idx} ({filename}) ---\n[OCR Error: {str(err)}]")
            
    return "\n\n".join(all_extracted)

def extract_text_from_pdf(pdf_path):
    """
    Extracts text page-by-page from a PDF file using PyMuPDF.
    If a page is an image-based or scanned PDF page (no digital text, or small header overlay with embedded images),
    it renders the page as a high-DPI image and invokes the OpenAI Vision model (gpt-4o-mini) for accurate OCR.
    """
    import fitz
    
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found at: {pdf_path}")
        
    doc = fitz.open(pdf_path)
    pages_text = []
    
    for page_idx, page in enumerate(doc):
        text = page.get_text().strip()
        
        # Calculate image area coverage on the page if images are present
        images = page.get_images(full=True)
        page_area = page.rect.width * page.rect.height if page.rect else 0
        img_area = 0
        if hasattr(page, 'get_image_info'):
            try:
                for info in page.get_image_info():
                    bbox = info.get('bbox')
                    if bbox:
                        img_area += (bbox[2] - bbox[0]) * (bbox[3] - bbox[1])
            except Exception:
                pass
        image_coverage = (img_area / page_area) if page_area > 0 else 0

        # Determine if Vision OCR should be invoked:
        # 1) Digital text is very short (< 50 chars)
        # 2) Page has embedded images and digital text is < 500 chars (scanned worksheet/form with a header)
        # 3) Image coverage is > 10% of page area and text length is < 800 chars
        # 4) Page has embedded images and low word count (< 40 words)
        words_count = len(text.split())
        should_run_ocr = (
            len(text) < 50 or
            (len(images) > 0 and len(text) < 500) or
            (image_coverage > 0.10 and len(text) < 800) or
            (len(images) > 0 and words_count < 40)
        )
        
        if should_run_ocr:
            logger.info(f"Page {page_idx + 1} detected as image-based/scanned (text len: {len(text)}, images: {len(images)}, coverage: {image_coverage:.1%}). Invoking OpenAI Vision OCR...")
            try:
                # Render page to PNG with 300 DPI for ultra-sharp text recognition
                pix = page.get_pixmap(dpi=300)
                png_bytes = pix.tobytes("png")
                base64_image = base64.b64encode(png_bytes).decode("utf-8")
                
                openai_client = get_openai_client()
                if openai_client:
                    response = openai_client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": (
                                    "You are an ultra-high precision Optical Character Recognition (OCR) model.\n"
                                    "Your task is to extract EVERY SINGLE piece of text visible on this document page image with 100% completeness and verbatim accuracy.\n\n"
                                    "STRICT OCR RULES:\n"
                                    "1. EXTRACT ALL VISIBLE TEXT: Include main body text, titles, headers, subheadings, footers, page numbers, questions, exercises, answers, table contents, cell text, bullet points, numbered lists, sidebars, logos/stamps with text, and fine print.\n"
                                    "2. DO NOT SUMMARIZE, SKIP, OR OMIT: Transcribe every single word, letter, number, punctuation mark, and symbol. Never skip text just because it is small, rotated, repeated, or formatted in columns/tables.\n"
                                    "3. DO NOT TRANSLATE: Output the extracted text in its exact original language and script.\n"
                                    "4. PRESERVE LAYOUT: Preserve paragraphs, table grid structures (using pipe | separators), line breaks, and indentation as much as possible.\n"
                                    "5. OUTPUT ONLY EXTRACTED TEXT: Return strictly the transcribed text with no introductory or concluding remarks."
                                )
                            },
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text", 
                                        "text": "Perform exhaustive 100% full-text OCR on this document page image. Extract every single word, header, footer, table entry, number, and note with absolute precision."
                                    },
                                    {
                                        "type": "image_url", 
                                        "image_url": {
                                            "url": f"data:image/png;base64,{base64_image}",
                                            "detail": "high"
                                        }
                                    }
                                ]
                            }
                        ],
                        max_tokens=4096,
                        temperature=0.0
                    )
                    ocr_text = response.choices[0].message.content.strip()
                    
                    if ocr_text:
                        # If OCR text extracted more text than digital text, or digital text was minimal
                        if len(ocr_text) > len(text) or len(text) < 50:
                            text = ocr_text
                        elif len(text) > 0 and ocr_text not in text:
                            text = f"{text}\n\n{ocr_text}"
                else:
                    logger.warning("OpenAI API client is not configured for Vision OCR fallback.")
            except Exception as ocr_err:
                logger.error(f"Failed Vision fallback OCR on page {page_idx + 1}: {str(ocr_err)}")
                
        pages_text.append(f"--- Page {page_idx + 1} ---\n{text}")
        
    try:
        doc.close()
    except Exception:
        pass
    return "\n\n".join(pages_text)

def extract_text_from_docx(docx_path):
    """
    Extracts text from a DOCX file using python-docx (including paragraphs and tables).
    """
    import docx
    
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found at: {docx_path}")
        
    doc = docx.Document(docx_path)
    lines = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
            
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                lines.append(" | ".join(row_text))
                
    return "\n".join(lines)

def extract_text_from_txt(txt_path):
    """
    Reads a raw text file cleanly, trying UTF-8 first, with error ignore fallback.
    """
    if not os.path.exists(txt_path):
        raise FileNotFoundError(f"TXT file not found at: {txt_path}")
        
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

def chunk_text(text, max_chars=3500):
    """
    Intelligently splits large text bodies into chunks not exceeding max_chars.
    Attempts to break along paragraph boundaries to maintain structural flow.
    """
    if not text:
        return []
        
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = []
    current_len = 0
    
    for p in paragraphs:
        # If adding paragraph exceeds chunk size limit
        if current_len + len(p) + 1 > max_chars:
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
                current_chunk = []
                current_len = 0
            
            # If a single paragraph is longer than max_chars, split by sentences
            if len(p) > max_chars:
                sentences = re.split(r'(?<=[.!?])\s+', p)
                sub_chunk = []
                sub_len = 0
                for s in sentences:
                    if sub_len + len(s) + 1 > max_chars:
                        if sub_chunk:
                            chunks.append(' '.join(sub_chunk))
                        sub_chunk = [s]
                        sub_len = len(s)
                    else:
                        sub_chunk.append(s)
                        sub_len += len(s) + 1
                if sub_chunk:
                    current_chunk = [' '.join(sub_chunk)]
                    current_len = sub_len
            else:
                current_chunk = [p]
                current_len = len(p)
        else:
            current_chunk.append(p)
            current_len += len(p) + 1
            
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
        
    return chunks

def translate_text_with_openai(text, target_lang, source_lang="auto"):
    """
    Translates the given text into the target language using OpenAI's API.
    Splits text into chunks, caches results, and handles retries / timeouts.
    """
    openai_client = get_openai_client()
    if not openai_client:
        raise ValueError("OpenAI API client is not configured. Please set OPENAI_API_KEY.")

    if not text or not text.strip():
        return ""

    chunks = chunk_text(text)
    translated_chunks = []

    for chunk in chunks:
        if not chunk.strip():
            translated_chunks.append("")
            continue

        # Check Cache to prevent redundant translation cost
        chunk_hash = hashlib.md5(f"{chunk}:{target_lang}:{source_lang}".encode("utf-8")).hexdigest()
        cache_key = f"openai_doc_chunk_{chunk_hash}"
        cached_translation = cache.get(cache_key)
        
        if cached_translation:
            translated_chunks.append(cached_translation)
            continue

        # Translation Prompts
        system_prompt = (
            "You are a professional, high-accuracy document translator. "
            "Translate the user input text completely into the requested target language. "
            "You must preserve original paragraph layouts, line breaks, item lists, tone, and names. "
            "Do not add any explanations, summaries, or intro/outro sentences. Output ONLY the translated text."
        )
        user_prompt = f"Source Language: {source_lang}\nTarget Language: {target_lang}\n\nText to translate:\n{chunk}"

        max_retries = 3
        backoff = 2
        translated_text = ""
        
        for attempt in range(max_retries):
            try:
                response = openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3,
                    timeout=45 # 45 seconds timeout
                )
                translated_text = response.choices[0].message.content.strip()
                # Cache the successful chunk translation for 1 day
                cache.set(cache_key, translated_text, timeout=86400)
                break
            except RateLimitError as e:
                if attempt == max_retries - 1:
                    logger.error(f"OpenAI RateLimit reached max retries: {str(e)}")
                    raise
                logger.warning(f"OpenAI RateLimit hit. Retrying in {backoff ** attempt} seconds...")
                time.sleep(backoff ** attempt)
            except (APIConnectionError, APIStatusError) as e:
                if attempt == max_retries - 1:
                    logger.error(f"OpenAI API connection/status failure reached max retries: {str(e)}")
                    raise
                logger.warning(f"OpenAI connection error: {str(e)}. Retrying in {backoff ** attempt} seconds...")
                time.sleep(backoff ** attempt)
            except Exception as e:
                logger.error(f"Unexpected error calling OpenAI API: {str(e)}")
                raise

        translated_chunks.append(translated_text)

    return "\n".join(translated_chunks)

def wrap_and_save_pdf(text_content, pdf_path, target_lang=None):
    """
    Utility to write text into a formatted multi-page PDF using PyMuPDF (fitz) Story API
    supporting proper Unicode font rendering for Indic and Latin scripts.
    """
    import fitz
    import html
    
    # Normalize target language code
    lang = (target_lang or '').lower().strip()
    
    # Map of language codes to Noto Sans font filenames
    font_mapping = {
        'ta': 'NotoSansTamil-Regular.ttf',
        'tamil': 'NotoSansTamil-Regular.ttf',
        'hi': 'NotoSansDevanagari-Regular.ttf',
        'hindi': 'NotoSansDevanagari-Regular.ttf',
        'mr': 'NotoSansDevanagari-Regular.ttf',
        'marathi': 'NotoSansDevanagari-Regular.ttf',
        'ne': 'NotoSansDevanagari-Regular.ttf',
        'nepali': 'NotoSansDevanagari-Regular.ttf',
        'te': 'NotoSansTelugu-Regular.ttf',
        'telugu': 'NotoSansTelugu-Regular.ttf',
        'ml': 'NotoSansMalayalam-Regular.ttf',
        'malayalam': 'NotoSansMalayalam-Regular.ttf',
        'kn': 'NotoSansKannada-Regular.ttf',
        'kannada': 'NotoSansKannada-Regular.ttf',
    }
    
    font_filename = font_mapping.get(lang, 'NotoSans-Regular.ttf')
    
    # Resolve the static fonts directory path
    fonts_dir = os.path.join(settings.BASE_DIR, 'static', 'fonts')
    font_file = os.path.join(fonts_dir, font_filename)
    
    # Fallback to NotoSans-Regular if the mapped font is missing
    if not os.path.exists(font_file):
        font_filename = 'NotoSans-Regular.ttf'
        font_file = os.path.join(fonts_dir, font_filename)
        
    font_family = font_filename.replace('.ttf', '').replace('-Regular', '')
    
    # Escape HTML and format paragraphs / line breaks
    escaped_text = html.escape(text_content)
    formatted_body = escaped_text.replace('\n', '<br/>')
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
    <style>
    @font-face {{
        font-family: '{font_family}';
        src: url('{font_filename}');
    }}
    body {{
        font-family: '{font_family}', sans-serif;
        font-size: 11pt;
        line-height: 1.5;
        color: #1a1a1a;
    }}
    </style>
    </head>
    <body>
    <div>{formatted_body}</div>
    </body>
    </html>
    """
    
    # Page dimensions: standard A4 (595 x 842 pt)
    mediabox = fitz.paper_rect("a4")
    # Margins (54 pt = 0.75 inch)
    where = mediabox + (54, 54, -54, -54)
    
    # Create the story with fonts_dir as archive source
    story = fitz.Story(html=html_content, archive=fonts_dir)
    writer = fitz.DocumentWriter(pdf_path)
    
    more = 1
    while more:
        device = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(device)
        writer.end_page()
        
    writer.close()

def save_as_docx(text, path):
    """
    Writes paragraphs into a DOCX file using python-docx.
    """
    import docx
    doc = docx.Document()
    for p in text.split('\n'):
        doc.add_paragraph(p)
    doc.save(path)

def generate_translated_file(translated_text, output_format, output_path, target_lang=None):
    """
    Generates a localized file based on output_format ('txt', 'pdf', 'docx').
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    if output_format == 'pdf':
        wrap_and_save_pdf(translated_text, output_path, target_lang)
    elif output_format == 'docx':
        save_as_docx(translated_text, output_path)
    else:
        # Default to txt
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(translated_text)
