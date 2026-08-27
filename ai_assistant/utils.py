import os
import io
import re
import uuid
import logging
import pypdf
import docx

logger = logging.getLogger(__name__)

# Max allowed resume upload size: 10MB
MAX_RESUME_SIZE_BYTES = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = ['.pdf', '.docx']


def extract_text_from_pdf(file_bytes_or_path):
    """
    Extracts text content from a PDF file using pypdf.
    Accepts either file bytes or file path.
    """
    try:
        if isinstance(file_bytes_or_path, (str, bytes)):
            if isinstance(file_bytes_or_path, str) and os.path.exists(file_bytes_or_path):
                reader = pypdf.PdfReader(file_bytes_or_path)
            else:
                stream = io.BytesIO(file_bytes_or_path) if isinstance(file_bytes_or_path, bytes) else file_bytes_or_path
                reader = pypdf.PdfReader(stream)
        else:
            reader = pypdf.PdfReader(file_bytes_or_path)

        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n".join(text_parts).strip()
        return full_text
    except Exception as e:
        logger.exception("Error extracting text from PDF")
        raise ValueError(f"Failed to read PDF resume: {str(e)}")


def extract_text_from_docx(file_bytes_or_path):
    """
    Extracts text content from a DOCX file using python-docx.
    Accepts either file bytes or file path.
    """
    try:
        if isinstance(file_bytes_or_path, bytes):
            stream = io.BytesIO(file_bytes_or_path)
            doc = docx.Document(stream)
        else:
            doc = docx.Document(file_bytes_or_path)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables as well
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts).strip()
        return full_text
    except Exception as e:
        logger.exception("Error extracting text from DOCX")
        raise ValueError(f"Failed to read DOCX resume: {str(e)}")


def extract_resume_text(uploaded_file):
    """
    Validates uploaded file extension & size, extracts text, and performs basic content validation.
    Returns (extracted_text, filename).
    """
    if not uploaded_file:
        raise ValueError("No resume file uploaded.")

    filename = uploaded_file.name
    ext = os.path.splitext(filename)[1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file format '{ext}'. Please upload a text-based PDF or DOCX file.")

    if uploaded_file.size > MAX_RESUME_SIZE_BYTES:
        raise ValueError("Resume file size exceeds maximum limit of 10MB.")

    # Read content bytes into memory
    file_bytes = uploaded_file.read()
    if not file_bytes:
        raise ValueError("The uploaded file is empty.")

    if ext == '.pdf':
        extracted_text = extract_text_from_pdf(file_bytes)
    elif ext == '.docx':
        extracted_text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported format.")

    # Validate extracted text content
    if not extracted_text or len(extracted_text.strip()) < 30:
        raise ValueError("Unable to extract text from this resume. Please upload a text-based PDF or DOCX file (scanned image PDFs without text layers are not supported).")

    return extracted_text.strip(), filename
